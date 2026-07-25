from __future__ import annotations

import io

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from ats_engine.generation.html_renderer import freeform_line_blocks, parse_freeform_document
from ats_engine.kit.contract import CoverLetterDocument, ResumeDocument

"""Deterministic rendering of structured Resume/Cover Letter documents into
ATS-friendly ``.docx`` bytes.

Mirrors ``generation/html_renderer.py``: single-column, no tables, no images,
no text boxes, no headers/footers — every element is a heading or a plain
paragraph/bullet so any ATS parser and a human reader see the same content in
the same order. ``python-docx`` is a pure-Python OOXML writer with no native
binary dependency, so — unlike WeasyPrint (confined to ``apps/api`` per
ADR-0004/ADR-0018) — this renderer lives in the engine, consistent with
``python-docx`` already being an engine dependency for input-side ``.docx``
parsing.
"""

_ACCENT = RGBColor(0x2F, 0x6F, 0x4F)
_INK = RGBColor(0x11, 0x11, 0x11)
_FONT_CLASSIC = "Georgia"
_FONT_MODERN = "Calibri"


def render_resume_docx(document: ResumeDocument, template: str) -> bytes:
    """Render a grounded, structured ``ResumeDocument`` into standalone DOCX bytes."""
    modern = template == "modern"
    doc = Document()
    _set_base_style(doc, modern)

    if document.candidate_name:
        _add_name_heading(doc, document.candidate_name, modern)
    if document.professional_headline:
        _add_paragraph(doc, document.professional_headline, bold=True, size=10.5)
    if document.contact_lines:
        _add_paragraph(doc, " | ".join(document.contact_lines), size=9.5, color=_muted(modern))

    if document.summary:
        _add_section_heading(doc, "Professional Summary", modern)
        _add_paragraph(doc, document.summary)

    if document.skill_groups:
        rows = [
            f"{group.label}: {', '.join(item for item in group.items if item)}"
            if group.label
            else ", ".join(item for item in group.items if item)
            for group in document.skill_groups
            if any(group.items)
        ]
        if rows:
            _add_section_heading(doc, "Technical Skills", modern)
            for row in rows:
                _add_paragraph(doc, row)

    if document.experience:
        _add_section_heading(doc, "Professional Experience", modern)
        for entry in document.experience:
            heading = " · ".join(part for part in (entry.employer, entry.location) if part)
            _add_entry_heading(doc, heading, entry.date_range)
            if entry.title:
                _add_paragraph(doc, entry.title, italic=True)
            for bullet in entry.bullets:
                if bullet:
                    _add_bullet(doc, bullet)

    if document.education:
        _add_section_heading(doc, "Education", modern)
        for edu_entry in document.education:
            heading = " · ".join(part for part in (edu_entry.institution, edu_entry.location) if part)
            _add_entry_heading(doc, heading, edu_entry.date_range)
            if edu_entry.degree:
                _add_paragraph(doc, edu_entry.degree, italic=True)
            for detail in edu_entry.details:
                if detail:
                    _add_paragraph(doc, detail)

    if document.certifications:
        certification_items = [
            " · ".join(part for part in (item.name, item.date, item.link) if part)
            for item in document.certifications
            if item.name
        ]
        if certification_items:
            _add_section_heading(doc, "Certifications", modern)
            for line in certification_items:
                _add_bullet(doc, line)

    for heading, lines in document.remaining_sections:
        content_lines = [line for line in lines if line]
        if content_lines:
            _add_section_heading(doc, heading or "Additional Information", modern)
            for line in content_lines:
                _add_paragraph(doc, line)

    return _to_bytes(doc)


def render_cover_letter_docx(document: CoverLetterDocument, template: str) -> bytes:
    """Render a grounded, structured ``CoverLetterDocument`` into standalone DOCX bytes."""
    modern = template == "modern"
    doc = Document()
    _set_base_style(doc, modern)

    if document.sender_name:
        _add_name_heading(doc, document.sender_name, modern, size=15)
    if document.sender_contact_lines:
        _add_paragraph(doc, " | ".join(document.sender_contact_lines), size=9.5, color=_muted(modern))
    if document.date:
        _add_paragraph(doc, document.date)

    recipient_lines = [
        line
        for line in (
            document.recipient_name,
            document.recipient_title,
            document.recipient_company,
            *document.recipient_address,
        )
        if line
    ]
    for line in recipient_lines:
        _add_paragraph(doc, line)

    if document.target_role:
        _add_paragraph(doc, document.target_role, bold=True)

    if document.greeting:
        _add_paragraph(doc, document.greeting)

    for paragraph in document.body_paragraphs:
        if paragraph:
            _add_paragraph(doc, paragraph)

    if document.closing:
        _add_paragraph(doc, document.closing)
    if document.signature_name:
        _add_paragraph(doc, document.signature_name)

    return _to_bytes(doc)


def render_plain_text_docx(text: str, *, template: str) -> bytes:
    """Render freeform text (a local edit with no structured document) as DOCX.

    Shares heading/bullet recognition with ``render_plain_text_html`` via
    :func:`ats_engine.generation.html_renderer.parse_freeform_document` /
    ``freeform_line_blocks`` so the PDF and DOCX exports of the same local edit
    agree; anything without a recognizable heading falls back to one paragraph
    per non-empty line rather than guessing at structure.
    """
    modern = template == "modern"
    doc = Document()
    _set_base_style(doc, modern)
    header_text, sections = parse_freeform_document(text)

    if not sections:
        for line in (text or "").splitlines():
            if line.strip():
                _add_paragraph(doc, line.strip())
        return _to_bytes(doc)

    if header_text:
        _add_paragraph(doc, header_text, size=9.5, color=_muted(modern))
    for heading, section_lines in sections:
        _add_section_heading(doc, heading, modern)
        for kind, line_text in freeform_line_blocks(section_lines):
            if kind == "bullet":
                _add_bullet(doc, line_text)
            else:
                _add_paragraph(doc, line_text)

    return _to_bytes(doc)


# --------------------------------------------------------------------------- #
# Shared helpers
# --------------------------------------------------------------------------- #
def _set_base_style(doc: DocxDocument, modern: bool) -> None:
    style = doc.styles["Normal"]
    style.font.name = _FONT_MODERN if modern else _FONT_CLASSIC
    style.font.size = Pt(10.5 if not modern else 10)
    style.font.color.rgb = _INK
    for section in doc.sections:
        section.top_margin = Pt(0.65 * 72)
        section.bottom_margin = Pt(0.65 * 72)
        section.left_margin = Pt(0.7 * 72)
        section.right_margin = Pt(0.7 * 72)


def _muted(modern: bool) -> RGBColor:
    return RGBColor(0x44, 0x44, 0x44) if modern else _INK


def _add_name_heading(doc: DocxDocument, name: str, modern: bool, *, size: float = 17) -> None:
    paragraph = doc.add_paragraph()
    if not modern:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name = name.upper()
    run = paragraph.add_run(name)
    run.bold = True
    run.font.size = Pt(size)
    if modern:
        run.font.color.rgb = _INK


def _add_paragraph(
    doc: DocxDocument,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 10.5,
    color: RGBColor | None = None,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _add_section_heading(doc: DocxDocument, title: str, modern: bool) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10 if modern else 9.5)
    if modern:
        run.font.color.rgb = _ACCENT
    _add_bottom_border(paragraph, color="2F6F4F" if modern else "111111")


def _add_entry_heading(doc: DocxDocument, heading_left: str, date_range: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    left = paragraph.add_run(heading_left)
    left.bold = True
    left.font.size = Pt(10.5)
    if date_range:
        right = paragraph.add_run(f"   {date_range}")
        right.font.size = Pt(10)


def _add_bullet(doc: DocxDocument, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)


def _add_bottom_border(paragraph: Paragraph, *, color: str) -> None:
    """Add a thin bottom border under a paragraph (section heading underline).

    A well-known ``python-docx`` recipe: OOXML has no high-level paragraph-
    border API, so this sets ``w:pBdr`` directly. Purely cosmetic — a border is
    not embedded text or an image, so it never affects ATS text extraction.
    """
    p = paragraph._p  # noqa: SLF001 - the documented low-level access point
    paragraph_properties = p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    paragraph_properties.append(border)


def _to_bytes(doc: DocxDocument) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


__all__ = ["render_cover_letter_docx", "render_plain_text_docx", "render_resume_docx"]
