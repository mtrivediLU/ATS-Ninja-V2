"""Delivery-first v7 acceptance and regression coverage.

All fixtures are synthetic and privacy-safe. They preserve the structural
features of the two production failure shapes without containing a real
candidate's identity, contact details, or credential identifiers.
"""

from __future__ import annotations

import io
import json
import logging
from copy import deepcopy
from functools import cache
from pathlib import Path
from typing import Any

import fitz
import pytest
from docx import Document as ReadDocument

import ats_engine.generation.optimizer as optimizer_module
from ats_engine import (
    ArtifactKind,
    ClaimStatus,
    ClaimType,
    DocumentState,
    KitState,
    application_kit_from_dict,
    application_kit_to_dict,
    generate_application_kit,
)
from ats_engine.config import EngineSettings
from ats_engine.evidence.resolver import resolve_requirements
from ats_engine.generation.docx_renderer import render_resume_docx
from ats_engine.generation.html_renderer import render_resume_html
from ats_engine.generation.latex_renderer import latex_escape
from ats_engine.generation.optimizer import (
    ResumeGateContext,
    build_resume_gate_context,
    optimize,
    validate_resume_plan_findings,
)
from ats_engine.generation.planning import (
    _build_headline,
    _build_summary,
    _headline,
    _rewrite_bullets_batch,
    build_resume_plan,
)
from ats_engine.generation.resume import generate_resume_text
from ats_engine.kit.change_actions import ChangeAction, apply_change_actions
from ats_engine.kit.contract import ApplicationKit
from ats_engine.models import (
    ContactInfo,
    EvidenceLink,
    Experience,
    JDProfile,
    PlacementAction,
    Profile,
    RequirementTerm,
    ResumePlan,
)
from ats_engine.parsing.document_extraction import extract_resume_document
from ats_engine.parsing.job_description import parse_jd
from ats_engine.parsing.resume import build_profile
from ats_engine.providers.base import LLMProvider
from ats_engine.scoring.ats_v2 import AtsScoreV2
from ats_engine.validation.calibration import apply_calibration, calibrate_identity
from ats_engine.validation.fidelity import (
    BulletPair,
    bullet_fidelity_findings,
    contains_fact,
    extract_named_entities,
    validate_resume_fidelity,
)
from ats_engine.validation.findings import CAL_FALSE_POSITIVE, ValidationSeverity
from ats_engine.validation.severity import (
    FIDELITY_MISSING_RESPONSIBILITY,
    FIDELITY_MISSING_TECHNOLOGY,
)
from ats_engine.validation.stuffing import validate_resume_stuffing_findings

FIXTURE_ROOT = Path(__file__).parent / "fixtures" / "real_cases"
RESUME_PATH = FIXTURE_ROOT / "resume" / "resume.txt"
CASE_TARGETS = {
    "coo_it_specialist": ("IT Specialist – Economic Development", "Chiefs of Ontario"),
    "claimsecure_it_admin": ("IT Administrator", "ClaimSecure"),
}
CASE_REQUIREMENT_SNAPSHOTS = {
    "coo_it_specialist": (
        "survey data",
        "geocoding",
        "information technology",
        "information management",
        "data analytics",
        "it support",
        "systems administration",
        "power bi desktop",
        "dax",
        "data modelling",
        "power query",
        "power bi service",
        "workspaces",
        "refresh cycles",
        "excel",
        "esri arcgis",
        "power bi",
        "data governance",
        "cybersecurity",
        "star schema",
        "troubleshooting",
        "multi-source data pipelines",
        "dashboards",
        "business requirements",
        "data integrity",
        "data ingestion",
        "data quality",
        "version control",
        "information architecture",
        "user interface design",
        "m-files",
        "document management",
        "ocap",
        "communication",
    ),
    "claimsecure_it_admin": (
        "active directory",
        "windows server",
        "microsoft 365",
        "troubleshooting",
        "networking",
        "dns",
        "dhcp",
        "vpn",
        "firewalls",
        "endpoint security",
        "vmware",
        "backups",
        "disaster recovery",
        "helpdesk",
        "virtualization",
        "patching",
        "asset inventory",
        "dashboards",
        "powershell",
        "python",
        "sql",
        "stakeholder communication",
    ),
}
EMPLOYERS = (
    "Northstar Health Analytics",
    "Cedar Ridge Systems",
    "Lakeview Municipal Services",
    "Meridian Software Labs",
    "Northern Data Research Institute",
    "Atlas Enterprise Consulting",
)
ROLE_FACTS = (
    "Senior Software Engineer",
    "Software Engineer",
    "Application Developer",
    "Data Analyst",
    "Systems Analyst",
)
DATE_FACTS = (
    "January 2022 - Present",
    "January 2020 - December 2021",
    "January 2018 - December 2019",
    "January 2016 - December 2017",
    "January 2014 - December 2015",
    "January 2012 - December 2013",
)
CREDENTIAL_IDS = (
    "TEST-PL300-001",
    "TEST-PL400-002",
    "TEST-AZ900-003",
    "TEST-SFDEV-004",
)
DELIVERED_STATES = {
    DocumentState.GENERATED,
    DocumentState.GENERATED_WITH_FALLBACK,
}


def _read_resume() -> str:
    return RESUME_PATH.read_text(encoding="utf-8")


def _read_jd(case: str) -> str:
    return (FIXTURE_ROOT / case / "job_description.txt").read_text(encoding="utf-8")


def _generate_case(case: str, *, prose_provider: LLMProvider | None = None) -> ApplicationKit:
    return generate_application_kit(
        resume_text=_read_resume(),
        job_description=_read_jd(case),
        include_resume=True,
        include_cover_letter=True,
        include_application_answers=False,
        include_job_fit=False,
        include_interview_prep=False,
        include_linkedin_outreach=False,
        settings=EngineSettings(
            tailoring_v2=True,
            delivery_first=True,
            llm_cache_enabled=False,
        ),
        use_llm=prose_provider is not None,
        prose_provider=prose_provider,
    )


@cache
def _cached_case(case: str) -> ApplicationKit:
    return _generate_case(case)


@pytest.mark.parametrize("case", tuple(CASE_TARGETS))
def test_real_case_target_and_requirement_snapshots_exclude_posting_noise(case: str) -> None:
    parsed = parse_jd(_read_jd(case))
    expected_title, expected_company = CASE_TARGETS[case]

    assert (parsed.title, parsed.company, parsed.parse_confidence) == (
        expected_title,
        expected_company,
        1.0,
    )
    assert tuple(requirement.canonical for requirement in parsed.requirements) == CASE_REQUIREMENT_SNAPSHOTS[case]
    requirement_text = " ".join(requirement.canonical for requirement in parsed.requirements).casefold()
    for forbidden in (
        "avery recruiter",
        "human resources",
        "opportunities@",
        "careers@",
        "salary",
        "references",
        "eligible to work",
        "inclusive recruitment",
    ):
        assert forbidden not in requirement_text


@pytest.mark.parametrize("case", tuple(CASE_TARGETS))
def test_real_cases_deliver_both_documents_with_monotone_ats_v2_scores(case: str) -> None:
    kit = _cached_case(case)
    expected_title, expected_company = CASE_TARGETS[case]

    assert kit.schema_version == "application-kit/v7"
    assert kit.state is KitState.COMPLETED
    assert (kit.target_role, kit.target_company, kit.target_confidence) == (
        expected_title,
        expected_company,
        1.0,
    )
    assert kit.validation.passed and not kit.validation.fatal
    assert kit.resume is not None and kit.resume.text and not kit.resume.validation.fatal
    assert kit.cover_letter is not None and kit.cover_letter.text and not kit.cover_letter.validation.fatal
    assert kit.delivery_reports[ArtifactKind.RESUME].state in DELIVERED_STATES
    assert kit.delivery_reports[ArtifactKind.COVER_LETTER].state in DELIVERED_STATES
    assert kit.match_report is not None
    assert kit.match_report.score_basis == "ats_v2"
    assert kit.match_report.tailored_ats_match is not None
    assert kit.match_report.tailored_ats_match.score >= kit.match_report.original_ats_match.score
    assert kit.match_report.optimization_trace.delivery_state in DELIVERED_STATES
    assert all(claim.status is ClaimStatus.SUPPORTED for claim in kit.all_claims())
    for artifact_text in (kit.resume.text, kit.cover_letter.text):
        assert "—" not in artifact_text
        assert "–" not in artifact_text
        assert "--" not in artifact_text


@pytest.mark.parametrize("case", tuple(CASE_TARGETS))
def test_real_cases_preserve_every_protected_resume_fact(case: str) -> None:
    source = _read_resume()
    profile = build_profile(source)
    kit = _cached_case(case)
    assert kit.resume is not None and kit.resume.document is not None
    delivered = kit.resume.text
    document = kit.resume.document

    for fact in (
        *EMPLOYERS,
        *ROLE_FACTS,
        *DATE_FACTS,
        *CREDENTIAL_IDS,
        "Toronto, ON (Remote)",
        "team of four engineers",
        "100% uptime",
        "SYNCONF 2023 Halifax",
        "Atlantic Innovation Centre",
        "Atlantic Technical University",
        "Journal of Fictional Applied Analytics",
        "Architected and built",
        "Streamlined Mission-Critical",
    ):
        assert fact in delivered
    for language in ("Python", "Java", "JavaScript", "SQL", "C", "C++", "PHP", "HTML5", "CSS3"):
        assert language in delivered
    assert document.remaining_sections == [
        (
            "Publications",
            [
                'Co-authored "Pattern Recognition for Synthetic Resource Models," '
                "Journal of Fictional Applied Analytics, 2023."
            ],
        )
    ]

    bullet_pairs = [
        BulletPair(
            original=source_bullet,
            candidate=document.experience[experience_index].bullets[bullet_index],
            location=f"experience:{experience_index}:bullet:{bullet_index}",
        )
        for experience_index, source_experience in enumerate(profile.experiences)
        for bullet_index, source_bullet in enumerate(source_experience.bullets)
    ]
    assert (
        validate_resume_fidelity(
            source,
            delivered,
            profile=profile,
            candidate_experiences=[
                Experience(
                    company=entry.employer,
                    title=entry.title,
                    location=entry.location,
                    dates=entry.date_range,
                    bullets=list(entry.bullets),
                )
                for entry in document.experience
            ],
            bullet_pairs=bullet_pairs,
        )
        == []
    )


@pytest.mark.parametrize("case", tuple(CASE_TARGETS))
def test_real_case_contract_is_byte_stable_without_a_provider(case: str) -> None:
    first = application_kit_to_dict(_cached_case(case))
    second = application_kit_to_dict(_generate_case(case))

    assert json.dumps(first, sort_keys=True, separators=(",", ":")) == json.dumps(
        second,
        sort_keys=True,
        separators=(",", ":"),
    )
    assert set(first["stage_timings"]["stages_ms"].values()) == {0}


def _render_resume_as_pdf(text: str) -> bytes:
    document = fitz.open()
    page: Any | None = None
    y = 0.0
    for line in text.splitlines():
        if page is None or y > 755:
            page = document.new_page(width=612, height=792)
            y = 36.0
        page.insert_text((30, y), line, fontname="helv", fontsize=7.5)
        y += 9.0
    content = document.tobytes()
    document.close()
    return content


def test_actual_pdf_extraction_preserves_real_case_structure_and_delivers() -> None:
    pdf_bytes = _render_resume_as_pdf(_read_resume())
    extracted = extract_resume_document(
        filename="synthetic-resume.pdf",
        content_type="application/pdf",
        content=pdf_bytes,
        max_bytes=10 * 1024 * 1024,
        max_pdf_pages=100,
        max_text_characters=100_000,
    )

    assert extracted.page_count == 1
    assert not extracted.warnings
    assert not extracted.manual_review_recommended
    for fact in (
        "Northstar Health Analytics",
        "SYNCONF 2023 Halifax",
        "TEST-PL300-001",
        "Streamlined Mission-Critical",
    ):
        assert fact in extracted.text

    profile = build_profile(extracted.text)
    assert len(profile.experiences) == 6
    kit = generate_application_kit(
        resume_text=extracted.text,
        job_description=_read_jd("coo_it_specialist"),
        include_resume=True,
        include_cover_letter=True,
        include_application_answers=False,
        include_job_fit=False,
        include_interview_prep=False,
        include_linkedin_outreach=False,
        settings=EngineSettings(delivery_first=True, llm_cache_enabled=False),
        use_llm=False,
    )
    assert kit.state is KitState.COMPLETED
    assert kit.resume is not None and kit.resume.text
    assert kit.cover_letter is not None and kit.cover_letter.text
    assert kit.match_report is not None and kit.match_report.tailored_ats_match is not None
    assert kit.match_report.tailored_ats_match.score >= kit.match_report.original_ats_match.score


@pytest.mark.parametrize("case", tuple(CASE_TARGETS))
def test_latex_docx_and_html_exports_retain_protected_facts_and_ats_safe_structure(case: str) -> None:
    kit = _cached_case(case)
    assert kit.resume is not None and kit.resume.document is not None

    latex = kit.resume.latex
    docx_bytes = render_resume_docx(kit.resume.document, "classic")
    docx = ReadDocument(io.BytesIO(docx_bytes))
    docx_text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    html = render_resume_html(kit.resume.document, "classic")

    assert latex and "\\begin{document}" in latex
    assert docx_bytes[:2] == b"PK"
    assert not docx.tables and not docx.inline_shapes
    assert "<table" not in html.casefold() and "<img" not in html.casefold()
    for fact in (
        *EMPLOYERS,
        *ROLE_FACTS,
        *DATE_FACTS,
        *CREDENTIAL_IDS,
        "team of four engineers",
        "100% uptime",
        "SYNCONF 2023 Halifax",
        "Journal of Fictional Applied Analytics",
    ):
        assert latex_escape(fact) in latex
        assert fact in docx_text
        assert fact in html


@pytest.mark.parametrize("case", tuple(CASE_TARGETS))
def test_real_cases_reject_restore_summary_is_stable_across_export_and_persistence(case: str) -> None:
    """Service-level regression coverage for the two production failure shapes,
    end to end: reject/restore of the composed summary must never withhold the
    résumé (the exact bug this branch's fix addresses), exports must keep
    working on the restored kit, protected facts and credential IDs must
    survive, and a persist/reload cycle must not destabilize the round trip.
    """
    kit = _generate_case(case)
    assert kit.resume is not None and kit.resume.document is not None
    assert kit.cover_letter is not None and kit.cover_letter.text
    delivered_summary = kit.resume.document.summary
    delivered_score = kit.match_report.tailored_ats_match.score
    assert delivered_summary

    resume_text = _read_resume()
    jd_text = _read_jd(case)

    rejected = apply_change_actions(
        kit=kit,
        resume_text=resume_text,
        job_description=jd_text,
        actions=[ChangeAction("resume::summary", "reject")],
        expected_revision=0,
    )
    assert rejected.ok, rejected.errors
    assert rejected.kit.resume is not None and rejected.kit.resume.document is not None
    assert rejected.kit.resume.document.summary != delivered_summary

    restored = apply_change_actions(
        kit=rejected.kit,
        resume_text=resume_text,
        job_description=jd_text,
        actions=[ChangeAction("resume::summary", "restore")],
        expected_revision=1,
    )
    # The exact regression this branch fixes: a validator false positive on the
    # composed summary (evidence-backed capability + certification wording +
    # the JD targeting clause) must never refuse the restore and accidentally
    # withhold the résumé.
    assert restored.ok, restored.errors
    assert restored.kit.resume is not None and restored.kit.resume.document is not None
    assert restored.kit.resume.text
    assert not restored.kit.resume.validation.fatal
    assert restored.kit.resume.document.summary == delivered_summary
    assert restored.kit.match_report is not None
    assert restored.kit.match_report.tailored_ats_match.score == delivered_score
    for credential_id in CREDENTIAL_IDS:
        assert credential_id in restored.kit.resume.text

    # Exports keep working on the restored kit (never a blanked artifact).
    docx_bytes = render_resume_docx(restored.kit.resume.document, "classic")
    assert docx_bytes[:2] == b"PK"
    pdf_bytes = _render_resume_as_pdf(restored.kit.resume.text)
    assert pdf_bytes[:5] == b"%PDF-"

    # Persist -> reload -> repeat the cycle: the immutable ledger baseline the
    # restore depends on must itself survive the real API/DB persistence path.
    persisted = json.loads(json.dumps(application_kit_to_dict(restored.kit)))
    reloaded = application_kit_from_dict(persisted)
    reloaded_rejected = apply_change_actions(
        kit=reloaded,
        resume_text=resume_text,
        job_description=jd_text,
        actions=[ChangeAction("resume::summary", "reject")],
        expected_revision=reloaded.revision,
    )
    assert reloaded_rejected.ok, reloaded_rejected.errors
    reloaded_restored = apply_change_actions(
        kit=reloaded_rejected.kit,
        resume_text=resume_text,
        job_description=jd_text,
        actions=[ChangeAction("resume::summary", "restore")],
        expected_revision=reloaded_rejected.kit.revision,
    )
    assert reloaded_restored.ok, reloaded_restored.errors
    assert reloaded_restored.kit.resume is not None and reloaded_restored.kit.resume.document is not None
    assert reloaded_restored.kit.resume.document.summary == delivered_summary


@pytest.fixture(scope="module")
def gated_coo_plan() -> tuple[Profile, JDProfile, ResumePlan, ResumeGateContext]:
    source = _read_resume()
    profile = build_profile(source)
    jd_profile = parse_jd(_read_jd("coo_it_specialist"), profile=profile, tailoring_v2=True)
    base_plan = build_resume_plan(
        contacts=profile.contact,
        jd_profile=jd_profile,
        profile=profile,
    )
    context = build_resume_gate_context(profile, base_plan)
    plan, _trace = optimize(
        profile,
        jd_profile,
        base_plan.requirements,
        base_plan.evidence_links,
        base_plan,
        gate_context=context,
    )
    return profile, jd_profile, plan, context


@pytest.mark.parametrize(
    ("deletion", "expected_fact"),
    [
        ("employer", "Northstar Health Analytics"),
        ("title", "Senior Software Engineer"),
        ("date", "January 2022 - Present"),
        ("location", "Toronto, ON (Remote)"),
        ("metric", "40%"),
        ("certification", "TEST-PL300-001"),
        ("technology", "PostgreSQL"),
        ("responsibility", "PostgreSQL"),
    ],
)
def test_same_calibrated_gate_context_keeps_genuine_fact_deletions_fatal(
    gated_coo_plan: tuple[Profile, JDProfile, ResumePlan, ResumeGateContext],
    deletion: str,
    expected_fact: str,
) -> None:
    profile, _jd_profile, source_plan, context = gated_coo_plan
    candidate = deepcopy(source_plan)
    if deletion == "employer":
        candidate.experience[0].company = ""
    elif deletion == "title":
        candidate.experience[0].title = ""
    elif deletion == "date":
        candidate.experience[0].dates = ""
    elif deletion == "location":
        candidate.experience[0].location = ""
    elif deletion == "metric":
        candidate.experience[0].bullets[1] = candidate.experience[0].bullets[1].replace("40%", "")
    elif deletion == "certification":
        candidate.certifications[0].credential_id = ""
    elif deletion == "technology":
        candidate.experience[0].bullets[0] = (
            candidate.experience[0]
            .bullets[0]
            .replace(
                "PostgreSQL",
                "database",
            )
        )
    elif deletion == "responsibility":
        candidate.experience[0].bullets[0] = "Architected and built a data warehouse."

    findings = validate_resume_plan_findings(
        candidate,
        profile,
        source_plan.requirements,
        candidate.placement_actions,
        context,
    )

    assert any(
        finding.severity is ValidationSeverity.FATAL and contains_fact(finding.fact, expected_fact)
        for finding in findings
    )


def test_authoritative_plan_gate_allows_experience_and_bullet_reordering(
    gated_coo_plan: tuple[Profile, JDProfile, ResumePlan, ResumeGateContext],
) -> None:
    profile, _jd_profile, source_plan, context = gated_coo_plan
    candidate = deepcopy(source_plan)
    candidate.experience = list(reversed(candidate.experience))
    for experience in candidate.experience:
        experience.bullets = list(reversed(experience.bullets))

    findings = validate_resume_plan_findings(
        candidate,
        profile,
        source_plan.requirements,
        candidate.placement_actions,
        context,
    )

    assert not any(finding.severity in {ValidationSeverity.FATAL, ValidationSeverity.DEGRADE} for finding in findings)


@pytest.mark.parametrize(
    ("source", "expected"),
    (
        (
            "Tata Consultancy Services - AICMSE: ZoomInfo; M-Files",
            ("Tata Consultancy Services", "AICMSE", "ZoomInfo", "M-Files"),
        ),
        (
            'Northwind Partners | "Acme Holdings" / Fabrikam Group',
            ("Northwind Partners", "Acme Holdings", "Fabrikam Group"),
        ),
        (
            "Power BI Service (Microsoft) – ESRI ArcGIS",
            ("ESRI", "ArcGIS"),
        ),
        ("Cloud SQL Auth Proxy; Google Cloud", ("Auth Proxy", "Google Cloud")),
        (
            "Tata Consultancy Services, Boston Consulting Group",
            ("Tata Consultancy Services", "Boston Consulting Group"),
        ),
        ("Acme Corp. Partnered with Beta Labs.", ("Acme Corp", "Beta Labs")),
        ("Acme, Inc.", ("Acme Inc",)),
        ("U.S. Bank collaborated with D3.js Labs.", ("U.S Bank", "D3.js Labs")),
    ),
)
def test_identity_projection_entities_are_self_contained_across_punctuation_corpus(
    source: str,
    expected: tuple[str, ...],
) -> None:
    entities = extract_named_entities(source)

    assert entities == expected
    assert bullet_fidelity_findings(source, source, source_text=source) == ()
    for entity in entities:
        assert contains_fact(source, entity)


def test_stuffing_calibration_identity_includes_exact_repetition_magnitude() -> None:
    source_term = "python python python python"
    five_term = next(
        finding
        for finding in validate_resume_stuffing_findings(
            "python python python python python",
            requirements=["python"],
            source_resume_text=source_term,
        )
        if finding.code == "STUFF_TERM_OCCURRENCES"
    )
    six_term = next(
        finding
        for finding in validate_resume_stuffing_findings(
            "python python python python python python",
            requirements=["python"],
            source_resume_text=source_term,
        )
        if finding.code == "STUFF_TERM_OCCURRENCES"
    )
    term_calibration = calibrate_identity([five_term])

    calibrated_five = apply_calibration([five_term], term_calibration)[0]
    calibrated_six = apply_calibration([six_term], term_calibration)[0]

    assert five_term.fact == "python | observed=5 | source-baseline=4"
    assert six_term.fact == "python | observed=6 | source-baseline=4"
    assert calibrated_five.code == CAL_FALSE_POSITIVE
    assert calibrated_five.severity is ValidationSeverity.WARN
    assert calibrated_six.code == "STUFF_TERM_OCCURRENCES"
    assert calibrated_six.severity is ValidationSeverity.DEGRADE

    source_bigram = "alpha beta alpha beta alpha beta"
    six_bigram = next(
        finding
        for finding in validate_resume_stuffing_findings(
            "alpha beta alpha beta alpha beta alpha beta alpha beta alpha beta",
            source_resume_text=source_bigram,
        )
        if finding.code == "STUFF_BIGRAM_REPETITION" and finding.fact.startswith("alpha beta |")
    )
    seven_bigram = next(
        finding
        for finding in validate_resume_stuffing_findings(
            "alpha beta alpha beta alpha beta alpha beta alpha beta alpha beta alpha beta",
            source_resume_text=source_bigram,
        )
        if finding.code == "STUFF_BIGRAM_REPETITION" and finding.fact.startswith("alpha beta |")
    )
    bigram_calibration = calibrate_identity([six_bigram])

    assert six_bigram.fact == "alpha beta | observed=6 | source-baseline=3"
    assert seven_bigram.fact == "alpha beta | observed=7 | source-baseline=3"
    assert apply_calibration([six_bigram], bigram_calibration)[0].code == CAL_FALSE_POSITIVE
    assert apply_calibration([seven_bigram], bigram_calibration)[0].code == "STUFF_BIGRAM_REPETITION"


def test_gate_calibration_uses_profile_identity_not_planner_content(
    gated_coo_plan: tuple[Profile, JDProfile, ResumePlan, ResumeGateContext],
) -> None:
    profile, _jd_profile, source_plan, _context = gated_coo_plan
    clean_context = build_resume_gate_context(profile, source_plan)
    contaminated = deepcopy(source_plan)
    contaminated.contacts.location = "Fictional Planner City"
    contaminated.role_identity = "Invented Executive"
    contaminated.headline = "Invented Executive | Unsupported Platform"
    contaminated.summary = "Unsupported planner summary."
    contaminated.skill_groups = [("Invented", ["Unsupported Platform"])]
    contaminated.experience = []
    contaminated.education = []
    contaminated.certifications = []
    contaminated.remaining_sections = []
    contaminated.placement_actions = []
    contaminated.plan_decisions = []

    contaminated_context = build_resume_gate_context(profile, contaminated)

    assert contaminated_context == clean_context

    deletion = deepcopy(source_plan)
    deletion.experience[0].title = ""
    findings = validate_resume_plan_findings(
        deletion,
        profile,
        source_plan.requirements,
        deletion.placement_actions,
        contaminated_context,
    )
    assert any(
        finding.severity is ValidationSeverity.FATAL and contains_fact(finding.fact, "Senior Software Engineer")
        for finding in findings
    )


def test_ordinary_technology_and_responsibility_deletions_are_fatal_but_semantic_rewrites_pass() -> None:
    technology = bullet_fidelity_findings(
        "Built resilient Airflow pipelines for analytics reporting and maintained operational runbooks.",
        "Built resilient pipelines for analytics reporting and maintained operational runbooks.",
        source_span="experience:0:bullet:0",
    )
    responsibility = bullet_fidelity_findings(
        "Owned incident response for production services and documented recurring issues.",
        "Owned production services and documented recurring issues.",
        source_span="experience:0:bullet:1",
    )
    semantic_rewrite = bullet_fidelity_findings(
        "Owned incident response for production services and documented recurring issues.",
        "Handled production incidents for services and documented recurring issues.",
        source_span="experience:0:bullet:1",
    )
    vocabulary_alias = bullet_fidelity_findings(
        "Built Apache Airflow pipelines.",
        "Built Airflow pipelines.",
        source_span="experience:0:bullet:2",
    )

    assert any(
        finding.code == FIDELITY_MISSING_TECHNOLOGY
        and finding.fact == "Airflow"
        and finding.severity is ValidationSeverity.FATAL
        and finding.source_span == "experience:0:bullet:0"
        for finding in technology
    )
    assert any(
        finding.code == FIDELITY_MISSING_RESPONSIBILITY
        and finding.fact == "incident response"
        and finding.severity is ValidationSeverity.FATAL
        and finding.source_span == "experience:0:bullet:1"
        for finding in responsibility
    )
    assert semantic_rewrite == ()
    assert vocabulary_alias == ()


def test_non_experience_bullet_headings_never_expand_raw_experience_scope() -> None:
    source = """Professional Experience
- Built Python reporting automation.
Projects
- Built an Airflow demonstration.
Publications
- Published an analytics article.
Awards
- Received a technical award.
Volunteering
- Supported a community event.
"""
    delivered = """Professional Experience
- Built Python reporting automation.
"""

    assert validate_resume_fidelity(source, delivered) == []


class _ResponseProvider:
    def __init__(self, response: str) -> None:
        self.response = response
        self.calls: list[str] = []

    @property
    def identity(self) -> str:
        return f"delivery-first-response:{id(self)}"

    def complete(self, prompt: str) -> str:
        self.calls.append(prompt)
        return self.response


class _RaisingProvider:
    def __init__(self, error: Exception) -> None:
        self.error = error
        self.calls = 0

    @property
    def identity(self) -> str:
        return f"delivery-first-raising:{id(self)}"

    def complete(self, _prompt: str) -> str:
        self.calls += 1
        raise self.error


def _structured_profile() -> Profile:
    return Profile(
        contact=ContactInfo(),
        retired_emails=[],
        role_identities=["Data Analyst"],
        tier_a={"python": "Python", "sql": "SQL"},
        tier_b={},
        tier_c={},
        adjacency={},
        experiences=[],
        education=[],
        certifications=[],
        supported_metrics=[],
        raw_markdown="Data Analyst with Python reporting experience.",
        source_summary="Data Analyst with Python reporting experience.",
    )


def _requirement(
    canonical: str,
    *,
    surface: str,
    kind: str,
    category: str,
) -> RequirementTerm:
    return RequirementTerm(
        canonical=canonical,
        surface=surface,
        aliases=(),
        kind=kind,
        section="required",
        weight=3.0,
        ngram=len(canonical.split()),
        category=category,
        jd_evidence_line=f"- {surface}",
    )


def _headline_links() -> list[EvidenceLink]:
    power_bi = _requirement(
        "power bi service",
        surface="Power BI Service",
        kind="tool",
        category="bi_analytics",
    )
    governance = _requirement(
        "data governance",
        surface="Data Governance",
        kind="methodology",
        category="security_governance",
    )
    return [
        EvidenceLink(
            requirement=power_bi,
            tier="A",
            resume_span="Built Power BI Service dashboards.",
            resume_location="experience:0:bullet:0",
            match_type="direct_experience",
            surface_to_use="Power BI Service",
            max_placement="headline",
        ),
        EvidenceLink(
            requirement=governance,
            tier="A",
            resume_span="Maintained Data Governance controls.",
            resume_location="experience:0:bullet:1",
            match_type="direct_experience",
            surface_to_use="Data Governance",
            max_placement="headline",
        ),
    ]


def test_structured_headline_success_uses_only_exact_credited_terms() -> None:
    links = _headline_links()
    fallback = _headline(JDProfile(), "Data Analyst", [], links)
    provider = _ResponseProvider(
        json.dumps(
            {
                "source_span": "resume:headline",
                "action": "rewrite_headline",
                "terms": ["Data Governance", "Power BI Service"],
                "text": "Data Analyst | Data Governance, Power BI Service",
            }
        )
    )

    result = _build_headline(
        fallback,
        "Data Analyst",
        JDProfile(title="BI Developer", company="Example"),
        links,
        provider,
    )

    assert result == "Data Analyst | Data Governance, Power BI Service"
    assert len(provider.calls) == 1
    assert "Exact source evidence" in provider.calls[0]
    assert "Explicit allowed terminology" in provider.calls[0]
    assert "Protected facts that cannot be removed or modified" in provider.calls[0]


@pytest.mark.parametrize(
    "response",
    (
        "not-json",
        json.dumps(
            {
                "source_span": "resume:headline",
                "action": "rewrite_headline",
                "terms": ["Kubernetes"],
                "text": "Data Analyst | Kubernetes",
            }
        ),
        json.dumps(
            {
                "source_span": "resume:headline",
                "action": "rewrite_headline",
                "terms": ["Power BI Service"],
                "text": "Director | Power BI Service",
            }
        ),
    ),
)
def test_structured_headline_malformed_unsupported_or_role_altered_output_falls_back(
    response: str,
) -> None:
    links = _headline_links()
    fallback = _headline(JDProfile(), "Data Analyst", [], links)

    assert (
        _build_headline(
            fallback,
            "Data Analyst",
            JDProfile(title="BI Developer", company="Example"),
            links,
            _ResponseProvider(response),
        )
        == fallback
    )


@pytest.mark.parametrize("error", (TimeoutError("timeout"), ConnectionError("outage")))
def test_structured_headline_timeout_or_outage_falls_back(error: Exception) -> None:
    links = _headline_links()
    fallback = _headline(JDProfile(), "Data Analyst", [], links)
    provider = _RaisingProvider(error)

    assert (
        _build_headline(
            fallback,
            "Data Analyst",
            JDProfile(title="BI Developer", company="Example"),
            links,
            provider,
        )
        == fallback
    )
    assert provider.calls == 1


def test_structured_summary_success_and_headline_failure_are_field_local() -> None:
    profile = _structured_profile()
    links = _headline_links()
    fallback_headline = _headline(JDProfile(), "Data Analyst", [], links)
    summary_text = (
        "Data Analyst with Python reporting experience. Builds reliable reports for finance teams, "
        "documents clear operating procedures, supports stakeholders with careful analysis, and keeps "
        "decisions grounded in verified source evidence for maintainable delivery."
    )

    class _MixedProvider:
        @property
        def identity(self) -> str:
            return f"delivery-first-mixed:{id(self)}"

        def complete(self, prompt: str) -> str:
            if "professional-summary rewrite" in prompt:
                return json.dumps(
                    {
                        "source_span": "resume:summary",
                        "action": "rewrite_summary",
                        "text": summary_text,
                    }
                )
            if "resume headline" in prompt:
                return json.dumps(
                    {
                        "source_span": "resume:headline",
                        "action": "rewrite_headline",
                        "terms": ["Kubernetes"],
                        "text": "Director | Kubernetes",
                    }
                )
            return ""

    provider = _MixedProvider()
    headline = _build_headline(
        fallback_headline,
        "Data Analyst",
        JDProfile(title="Data Analyst", company="Example"),
        links,
        provider,
    )
    summary = _build_summary(
        "Data Analyst",
        ["Python"],
        JDProfile(title="Data Analyst", company="Example"),
        profile,
        None,
        links,
        provider,
    )

    assert headline == fallback_headline
    assert summary == summary_text


def _bullet_inputs() -> tuple[list[str], Profile, JDProfile]:
    bullets = [
        "Built Python reporting automation for finance users.",
        "Documented SQL release procedures for operations.",
    ]
    return (
        bullets,
        _structured_profile(),
        JDProfile(title="Data Analyst", company="Example"),
    )


def test_structured_bullet_rewrites_succeed_and_fall_back_per_item() -> None:
    bullets, profile, jd_profile = _bullet_inputs()
    provider = _ResponseProvider(
        json.dumps(
            [
                {
                    "source_span": "resume::bullet0",
                    "action": "rewrite_bullet",
                    "text": "Created Python reporting automation for finance users.",
                },
                {
                    "source_span": "resume::bullet1",
                    "action": "wrong_action",
                    "text": "Prepared SQL release procedures for operations.",
                },
            ]
        )
    )

    rewritten = _rewrite_bullets_batch(
        bullets,
        jd_profile,
        ["python", "sql"],
        profile,
        provider,
    )

    assert rewritten == [
        "Created Python reporting automation for finance users.",
        bullets[1],
    ]
    assert len(provider.calls) == 1
    for contract_field in (
        "Structured source inputs",
        "source_evidence",
        "allowed_terminology",
        "protected_facts",
    ):
        assert contract_field in provider.calls[0]


def test_malformed_bullet_batch_falls_back_without_cross_item_loss() -> None:
    bullets, profile, jd_profile = _bullet_inputs()
    provider = _ResponseProvider("not-json")

    rewritten = _rewrite_bullets_batch(
        bullets,
        jd_profile,
        ["python", "sql"],
        profile,
        provider,
    )

    assert rewritten == bullets
    assert len(provider.calls) == 2


@pytest.mark.parametrize(
    "error_type",
    (TimeoutError, ConnectionError),
)
def test_bullet_timeout_or_outage_is_private_and_delivers_source_fallback(
    error_type: type[Exception],
    caplog: pytest.LogCaptureFixture,
) -> None:
    bullets, profile, jd_profile = _bullet_inputs()
    secret = "PRIVATE-CANDIDATE-CONTENT api-key-example"
    provider = _RaisingProvider(error_type(secret))

    with caplog.at_level(logging.WARNING):
        rewritten = _rewrite_bullets_batch(
            bullets,
            jd_profile,
            ["python", "sql"],
            profile,
            provider,
        )

    assert rewritten == bullets
    assert provider.calls == 1
    assert secret not in caplog.text
    assert error_type.__name__ in caplog.text


def test_valid_structured_summary_is_truth_gated_without_losing_safe_placements(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fabricated_management = "Managed an engineering department"
    proposed_summary = (
        "Business Intelligence Developer and Senior Software Engineer experienced in reliable data products, "
        "data warehousing, analytics, and cloud services. Managed an engineering department and owned executive "
        "hiring decisions while delivering careful operational reporting for stakeholders across the organization."
    )

    class _FabricatedSummaryProvider:
        @property
        def identity(self) -> str:
            return "delivery-first-fabricated-summary"

        def complete(self, prompt: str) -> str:
            if "professional-summary rewrite" in prompt:
                return json.dumps(
                    {
                        "source_span": "resume:summary",
                        "action": "rewrite_summary",
                        "text": proposed_summary,
                    }
                )
            return "{}"

    # Isolate the placement loop from the independent deterministic summary
    # quality proposal. With that proposal held at the source floor, the same
    # optimizer run must reject unsafe AI prose and still accept safe,
    # score-improving PlacementAction values.
    source_summary = (
        "Data Analyst experienced in reporting automation, careful analysis, stakeholder support, "
        "operational documentation, and maintainable delivery."
    )
    source_bullet = (
        "Built Python reporting automation for finance stakeholders with clear documentation "
        "and reliable scheduled delivery."
    )
    source_resume = f"""{source_summary}
Professional Experience
Synthetic Analytics | Remote
Data Analyst
2022 - Present
- {source_bullet}
Technical Skills
Languages: SQL
"""
    profile = Profile(
        contact=ContactInfo(location="Remote"),
        retired_emails=[],
        role_identities=["Data Analyst"],
        tier_a={"python": "Python"},
        tier_b={},
        tier_c={"sql": "SQL"},
        adjacency={},
        experiences=[
            Experience(
                company="Synthetic Analytics",
                title="Data Analyst",
                location="Remote",
                dates="2022 - Present",
                bullets=[source_bullet],
            )
        ],
        education=[],
        certifications=[],
        supported_metrics=[],
        raw_markdown=source_resume,
        source_summary=source_summary,
        source_skill_groups=[("Languages", ["SQL"])],
    )
    jd_profile = JDProfile(
        title="Platform Data Analyst",
        company="Example",
        required_qualifications=["Python", "Kubernetes"],
        technical_keywords=["Python", "Kubernetes"],
        requirements=[
            _requirement(
                "python",
                surface="Python",
                kind="language",
                category="programming",
            ),
            _requirement(
                "kubernetes",
                surface="Kubernetes",
                kind="platform",
                category="platform",
            ),
        ],
    )
    base_plan = build_resume_plan(
        contacts=profile.contact,
        jd_profile=jd_profile,
        profile=profile,
    )
    floor = optimizer_module._source_content_plan(
        base_plan,
        profile,
        base_plan.evidence_links,
    )
    attack_plan = deepcopy(base_plan)
    attack_plan.summary = (
        f"{source_summary} Managed an engineering department and owned an engineering roadmap while delivering "
        "careful analysis, clear operational documentation, reliable stakeholder communication, practical "
        "reporting support, and maintainable work across the organization."
    )
    with monkeypatch.context() as patch:
        patch.setattr(
            optimizer_module,
            "rewrite_summary",
            lambda *_args, **_kwargs: floor.summary,
        )
        _delivered, direct_trace = optimize(
            profile,
            jd_profile,
            attack_plan.requirements,
            attack_plan.evidence_links,
            attack_plan,
            gate_context=build_resume_gate_context(profile, attack_plan),
            accept_generated_prose=True,
        )

    assert "ai:summary" not in direct_trace.accepted_actions
    assert any(rejection.action == "ai:summary" for rejection in direct_trace.rejected_actions)
    assert any(
        action.startswith(("mention_summary:", "append_skill:", "surface_variant:", "weave_bullet:"))
        for action in direct_trace.accepted_actions
    )

    kit = _generate_case(
        "coo_it_specialist",
        prose_provider=_FabricatedSummaryProvider(),
    )

    assert kit.resume is not None and kit.resume.text
    assert kit.match_report is not None
    trace = kit.match_report.optimization_trace
    assert "ai:summary" not in trace.accepted_actions
    assert any(rejection.action == "ai:summary" for rejection in trace.rejected_actions)
    assert fabricated_management.casefold() not in kit.resume.text.casefold()
    assert any(
        claim.claim_type is ClaimType.MANAGEMENT
        and claim.status is ClaimStatus.REPAIRED
        and fabricated_management.casefold() in claim.text.casefold()
        for claim in kit.resume.claims
    )
    assert not kit.resume.validation.fatal


def test_accepted_bridge_is_explicit_in_reviewable_plan_decision(
    gated_coo_plan: tuple[Profile, JDProfile, ResumePlan, ResumeGateContext],
) -> None:
    profile, _jd_profile, source_plan, _context = gated_coo_plan
    bridge_bullets = [
        "Maintained ETL/ELT jobs for scheduled reporting.",
        "Built ingestion pipelines from Salesforce and HubSpot for finance reporting.",
    ]
    bridge_profile = _structured_profile()
    bridge_profile.experiences = [
        Experience(
            company="Synthetic Data Systems",
            title="Data Engineer",
            location="Remote",
            dates="2022 - 2025",
            bullets=bridge_bullets,
        )
    ]
    bridge_profile.raw_markdown = "\n".join(bridge_bullets)
    link = resolve_requirements(
        [
            _requirement(
                "multi-source data pipelines",
                surface="multi-source data pipelines",
                kind="skill",
                category="data_engineering",
            )
        ],
        bridge_profile,
        bridge_profile.raw_markdown,
    )[0]
    assert (
        link.tier,
        link.match_type,
        link.supporting_spans,
        link.supporting_locations,
    ) == (
        "A",
        "bridged",
        tuple(bridge_bullets),
        ("experience:0:bullet:0", "experience:0:bullet:1"),
    )
    action = PlacementAction(
        term=link.surface_to_use or link.requirement.surface,
        link=link,
        target="summary",
        operation="mention_summary",
        rendered_text=link.surface_to_use or link.requirement.surface,
        grounded_by=" | ".join(link.supporting_locations),
    )
    candidate = optimizer_module._apply_actions(source_plan, [action], profile)

    decisions = optimizer_module._v2_plan_decisions(
        [],
        candidate,
        profile,
        [action],
    )
    summary_decision = next(decision for decision in decisions if decision.location_id == "resume::summary")

    assert action.term in summary_decision.matched_keywords
    assert "Conservative evidence bridge:" in summary_decision.reason
    assert f"{action.term} ({len(link.supporting_spans)} candidate source spans)" in summary_decision.reason
    assert action.grounded_by == " | ".join(link.supporting_locations)


def test_failed_action_batch_bisects_and_keeps_safe_sibling(
    gated_coo_plan: tuple[Profile, JDProfile, ResumePlan, ResumeGateContext],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    profile, _jd_profile, source_plan, context = gated_coo_plan
    link = next(link for link in source_plan.evidence_links if link.tier != "missing")
    safe = PlacementAction(
        term=link.surface_to_use or link.requirement.surface,
        link=link,
        target="summary",
        operation="mention_summary",
        rendered_text=link.surface_to_use or link.requirement.surface,
        grounded_by=link.resume_location,
    )
    unsafe = PlacementAction(
        term=link.surface_to_use or link.requirement.surface,
        link=link,
        target="unsafe",
        operation="unsafe_delete",
        rendered_text="",
        grounded_by=link.resume_location,
    )

    def corrupt_only_unsafe(
        base: ResumePlan,
        actions: list[PlacementAction],
        _profile: Profile,
    ) -> ResumePlan:
        candidate = deepcopy(base)
        candidate.placement_actions = [*base.placement_actions, *actions]
        if any(action.target == "unsafe" for action in actions):
            candidate.experience[0].company = ""
        return candidate

    monkeypatch.setattr(optimizer_module, "_apply_actions", corrupt_only_unsafe)
    plan, actions, rejected = optimizer_module._accept_safe_actions(
        source_plan,
        [],
        [safe, unsafe],
        profile,
        source_plan.requirements,
        source_plan.evidence_links,
        profile.raw_markdown,
        context,
    )

    assert actions == [safe]
    assert plan.experience[0].company == profile.experiences[0].company
    assert [item.action for item in rejected] == [f"{unsafe.operation}:{unsafe.target}:{unsafe.term}"]


def test_all_unsafe_actions_deliver_validated_source_floor(
    gated_coo_plan: tuple[Profile, JDProfile, ResumePlan, ResumeGateContext],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    profile, jd_profile, base_plan, context = gated_coo_plan
    base_plan = deepcopy(base_plan)
    base_plan.headline = base_plan.role_identity
    link = next(link for link in base_plan.evidence_links if link.tier != "missing")
    unsafe = PlacementAction(
        term=link.surface_to_use or link.requirement.surface,
        link=link,
        target="unsafe",
        operation="unsafe_delete",
        rendered_text="",
        grounded_by=link.resume_location,
    )

    def corrupt_actions(
        base: ResumePlan,
        actions: list[PlacementAction],
        _profile: Profile,
    ) -> ResumePlan:
        candidate = deepcopy(base)
        if actions:
            candidate.experience[0].company = ""
        return candidate

    target_title = jd_profile.title.replace("–", "-").replace("—", "-")
    source_summary = f"{profile.source_summary} Targeting {target_title} opportunities."
    monkeypatch.setattr(optimizer_module, "plan_placements", lambda *_args: [unsafe])
    monkeypatch.setattr(optimizer_module, "_apply_actions", corrupt_actions)
    monkeypatch.setattr(
        optimizer_module,
        "rewrite_summary",
        lambda *_args, **_kwargs: source_summary,
    )
    monkeypatch.setattr(
        optimizer_module,
        "score_resume_v2",
        lambda *_args, **_kwargs: AtsScoreV2(
            score=50.0,
            base_score=50.0,
            density_penalty=0.0,
            placement_bonus=0.0,
        ),
    )

    delivered, trace = optimize(
        profile,
        jd_profile,
        base_plan.requirements,
        base_plan.evidence_links,
        base_plan,
        gate_context=context,
    )

    assert trace.delivery_state is DocumentState.GENERATED_WITH_FALLBACK
    assert trace.fallback_reason
    assert not trace.accepted_actions
    assert trace.rejected_actions
    assert [experience.company for experience in delivered.experience] == [
        experience.company for experience in profile.experiences
    ]
    assert not any(
        finding.severity in {ValidationSeverity.FATAL, ValidationSeverity.DEGRADE}
        for finding in validate_resume_plan_findings(
            delivered,
            profile,
            delivered.requirements,
            delivered.placement_actions,
            context,
        )
    )
    assert generate_resume_text(delivered)
