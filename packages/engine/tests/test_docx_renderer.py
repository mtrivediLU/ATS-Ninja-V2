from __future__ import annotations

import io

from docx import Document as ReadDocument

from ats_engine.generation.docx_renderer import (
    render_cover_letter_docx,
    render_plain_text_docx,
    render_resume_docx,
)
from ats_engine.kit.contract import (
    CoverLetterDocument,
    ResumeCertificationEntry,
    ResumeDocument,
    ResumeEducationEntry,
    ResumeExperienceEntry,
    ResumeSkillGroup,
)

"""DOCX export renders ATS-friendly, single-column .docx bytes for both the
Resume and Cover Letter, mirroring the existing HTML/PDF renderer's structure
(no tables, no images, no text boxes) so a human reader and an ATS text
extractor see the same content in the same order.
"""


def _resume_document() -> ResumeDocument:
    return ResumeDocument(
        candidate_name="Jordan Lee",
        professional_headline="Backend Engineer",
        contact_lines=["jordan@example.com", "linkedin.com/in/jordanlee"],
        summary="Backend engineer building Python services.",
        skill_groups=[ResumeSkillGroup("Programming Languages", ["Python", "SQL"])],
        experience=[
            ResumeExperienceEntry(
                employer="Acme Corp",
                title="Software Engineer",
                location="Remote",
                date_range="2019 - 2024",
                bullets=["Built Python services and REST APIs", "Documented release procedures"],
            )
        ],
        education=[
            ResumeEducationEntry(
                institution="State University",
                degree="Bachelor of Computer Science",
                location="",
                date_range="2015 - 2019",
                details=[],
            )
        ],
        certifications=[ResumeCertificationEntry("AWS Certified Developer", "2022", "")],
        remaining_sections=[("Projects", ["Built an internal tool"])],
    )


def _cover_letter_document() -> CoverLetterDocument:
    return CoverLetterDocument(
        sender_name="Jordan Lee",
        sender_contact_lines=["jordan@example.com"],
        date="2026-07-25",
        recipient_name="Hiring Manager",
        recipient_title="",
        recipient_company="Acme Corp",
        recipient_address=[],
        target_role="Backend Engineer",
        greeting="Dear Hiring Manager,",
        body_paragraphs=["First paragraph about fit.", "Second paragraph about experience."],
        closing="Sincerely,",
        signature_name="Jordan Lee",
    )


def _paragraph_texts(docx_bytes: bytes) -> list[str]:
    read = ReadDocument(io.BytesIO(docx_bytes))
    return [p.text for p in read.paragraphs]


def test_resume_docx_renders_all_sections_in_order() -> None:
    docx_bytes = render_resume_docx(_resume_document(), "classic")
    assert docx_bytes[:2] == b"PK"  # a real OOXML (zip) package, not empty bytes
    texts = _paragraph_texts(docx_bytes)
    joined = "\n".join(texts)

    # Candidate identity and every structured section is present.
    assert "JORDAN LEE" in joined  # classic template uppercases the name
    assert "Backend Engineer" in joined
    assert "jordan@example.com" in joined
    assert "Backend engineer building Python services." in joined
    assert "Python" in joined and "SQL" in joined
    assert "Built Python services and REST APIs" in joined
    assert "Documented release procedures" in joined
    assert "Acme Corp" in joined
    assert "2019 - 2024" in joined
    assert "State University" in joined
    assert "AWS Certified Developer" in joined
    assert "Built an internal tool" in joined

    # Order: name/summary precede experience, which precedes education.
    assert texts.index("Backend engineer building Python services.") < next(
        i for i, t in enumerate(texts) if "Acme Corp" in t
    )
    assert next(i for i, t in enumerate(texts) if "Acme Corp" in t) < next(
        i for i, t in enumerate(texts) if "State University" in t
    )


def test_resume_docx_modern_template_does_not_uppercase_name() -> None:
    docx_bytes = render_resume_docx(_resume_document(), "modern")
    texts = _paragraph_texts(docx_bytes)
    assert "Jordan Lee" in texts


def test_resume_docx_has_no_tables_or_images() -> None:
    # ATS-safety: single-column text only, no tables/images that could scramble
    # extraction order.
    docx_bytes = render_resume_docx(_resume_document(), "classic")
    read = ReadDocument(io.BytesIO(docx_bytes))
    assert len(read.tables) == 0
    assert len(read.inline_shapes) == 0


def test_resume_docx_omits_empty_sections() -> None:
    minimal = ResumeDocument(candidate_name="Sam Rivera")
    docx_bytes = render_resume_docx(minimal, "classic")
    texts = _paragraph_texts(docx_bytes)
    assert texts == ["SAM RIVERA"]


def test_cover_letter_docx_renders_all_parts_in_order() -> None:
    docx_bytes = render_cover_letter_docx(_cover_letter_document(), "classic")
    texts = _paragraph_texts(docx_bytes)
    joined = "\n".join(texts)
    assert "JORDAN LEE" in joined
    assert "jordan@example.com" in joined
    assert "Hiring Manager" in joined
    assert "Acme Corp" in joined
    assert "Backend Engineer" in joined
    assert "Dear Hiring Manager," in joined
    assert "First paragraph about fit." in joined
    assert "Second paragraph about experience." in joined
    assert "Sincerely," in joined

    greeting_index = texts.index("Dear Hiring Manager,")
    first_paragraph_index = texts.index("First paragraph about fit.")
    closing_index = texts.index("Sincerely,")
    assert greeting_index < first_paragraph_index < closing_index


def test_docx_renderer_is_deterministic() -> None:
    doc = _resume_document()
    first = render_resume_docx(doc, "classic")
    second = render_resume_docx(doc, "classic")
    assert _paragraph_texts(first) == _paragraph_texts(second)


def test_plain_text_docx_recognizes_headings_when_present() -> None:
    text = "Jordan Lee\njordan@example.com\n\nExperience\n- Built Python services\n\nSkills\nPython, SQL\n"
    docx_bytes = render_plain_text_docx(text, template="classic")
    texts = _paragraph_texts(docx_bytes)
    joined = "\n".join(texts)
    assert "Jordan Lee" in joined
    assert "EXPERIENCE" in joined
    assert "Built Python services" in joined
    assert "SKILLS" in joined
    assert "Python, SQL" in joined


def test_plain_text_docx_falls_back_to_verbatim_without_recognized_headings() -> None:
    text = "Just some freeform text\nwith no recognizable heading at all"
    docx_bytes = render_plain_text_docx(text, template="classic")
    texts = _paragraph_texts(docx_bytes)
    assert texts == ["Just some freeform text", "with no recognizable heading at all"]
