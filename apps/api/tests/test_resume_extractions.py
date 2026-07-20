from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject


def _text_pdf(*pages: str) -> bytes:
    writer = PdfWriter()
    for text in pages:
        page = writer.add_blank_page(width=612, height=792)
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        font_ref = writer._add_object(font)
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
        )
        stream = DecodedStreamObject()
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode())
        page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _docx() -> bytes:
    document = Document()
    document.add_heading("Synthetic Resume", level=1)
    document.add_paragraph("Synthetic summary for document extraction.")
    document.add_paragraph("Built local systems", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.asyncio
async def test_extracts_txt_and_returns_only_safe_metadata(client: object) -> None:
    response = await client.post(
        "/api/v1/resume-extractions",
        files={"file": ("resume.txt", b"Synthetic Candidate\nBuilt a local analytics service.", "text/plain")},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["extraction_method"] == "plain_text"
    assert body["text"] == "Synthetic Candidate\nBuilt a local analytics service."
    assert body["page_count"] is None
    assert "path" not in body
    assert body["truncated"] is False


@pytest.mark.asyncio
async def test_extracts_pdf_in_page_order(client: object) -> None:
    response = await client.post(
        "/api/v1/resume-extractions",
        files={"file": ("resume.pdf", _text_pdf("First synthetic page", "Second synthetic page"), "application/pdf")},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["extraction_method"] == "pdf_text"
    assert body["page_count"] == 2
    assert body["text"].index("First synthetic page") < body["text"].index("Second synthetic page")


@pytest.mark.asyncio
async def test_extracts_pdf_bullet_glued_to_text_gets_a_restored_space(client: object) -> None:
    """PDF text extraction commonly places a bullet glyph directly against its
    text (no literal space codepoint, since the visual gap is glyph
    positioning). Regression for the false Resume-withholding defect: the
    extracted, reviewable text must read as a normal bullet line."""
    response = await client.post(
        "/api/v1/resume-extractions",
        files={
            "file": (
                "resume.pdf",
                _text_pdf("*Managed cloud infrastructure across two regions."),
                "application/pdf",
            )
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert "* Managed cloud infrastructure across two regions." in body["text"]


@pytest.mark.asyncio
async def test_extracts_docx_paragraphs_lists_and_tables(client: object) -> None:
    response = await client.post(
        "/api/v1/resume-extractions",
        files={
            "file": (
                "resume.docx",
                _docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["extraction_method"] == "docx_text"
    assert "Synthetic Resume" in body["text"]
    assert "- Built local systems" in body["text"]
    assert "Skill | Python" in body["text"]


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("filename", "content", "content_type", "expected"),
    [
        ("resume.doc", b"legacy", "application/msword", "Legacy .doc files are not supported yet."),
        ("resume.docm", b"zip", "application/vnd.ms-word.document.macroEnabled.12", "Upload a PDF, DOCX, or TXT"),
        ("resume.pdf", b"not a pdf", "application/pdf", "malformed"),
        ("resume.txt", b"\x00\x01\x02\x03", "text/plain", "binary data"),
        ("../resume.txt", b"Synthetic resume text long enough.", "text/plain", "Upload a PDF, DOCX, or TXT"),
    ],
)
async def test_rejects_unsafe_or_unsupported_uploads(
    client: object,
    filename: str,
    content: bytes,
    content_type: str,
    expected: str,
) -> None:
    response = await client.post("/api/v1/resume-extractions", files={"file": (filename, content, content_type)})

    assert response.status_code == 422
    assert expected in response.json()["detail"]
    assert "Traceback" not in response.text


@pytest.mark.asyncio
async def test_rejects_scanned_pdf_and_mime_spoofing(client: object) -> None:
    scanned = await client.post(
        "/api/v1/resume-extractions",
        files={"file": ("resume.pdf", _text_pdf(""), "application/pdf")},
    )
    spoofed = await client.post(
        "/api/v1/resume-extractions",
        files={"file": ("resume.pdf", _text_pdf("Synthetic readable resume text"), "text/plain")},
    )

    assert scanned.status_code == 422
    assert "No readable text was found" in scanned.json()["detail"]
    assert spoofed.status_code == 422
    assert "does not match" in spoofed.json()["detail"]
